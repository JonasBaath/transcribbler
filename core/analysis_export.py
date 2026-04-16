"""
analysis_export.py — Export analysis view excerpts in various formats.
"""
import csv
import io
from .export import _code_path


def _group_excerpts(excerpts, mode):
    """Group excerpts by code. Returns [{code_id, code_name, code_color, code_number, code_path, excerpts: [...]}]."""
    groups = {}
    order = []
    for e in excerpts:
        cid = e["code_id"]
        if cid not in groups:
            groups[cid] = {
                "code_id": cid,
                "code_name": e["code_name"],
                "code_color": e["code_color"],
                "code_number": e.get("code_number", ""),
                "code_path": e.get("code_path", ""),
                "excerpts": [],
            }
            order.append(cid)
        groups[cid]["excerpts"].append(e)
    return [groups[cid] for cid in order]


def export_analysis_md(project, excerpts, mode="separate"):
    """Export excerpts as Markdown."""
    lines = [f"# Analys — {project['name']}\n"]
    grouped = _group_excerpts(excerpts, mode)
    for g in grouped:
        prefix = f"{g['code_number']}. " if g["code_number"] else ""
        lines.append(f"\n## {prefix}{g['code_path'] or g['code_name']}\n")
        for e in g["excerpts"]:
            label = e.get("transcript_label", "")
            name = e.get("transcript_name", "")
            anchor = " 📌" if e.get("anchor") else ""
            lines.append(f"**{label}. {name}** _(kodare: {e['coder']})_{anchor}")
            lines.append(f"> {e['text']}\n")
            if e.get("memo"):
                lines.append(f"> _Memo: {e['memo']}_\n")
    return "\n".join(lines)


def export_analysis_csv(project, excerpts, mode="separate"):
    """Export excerpts as CSV."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "code_number", "code_name", "code_path", "transcript_label",
        "transcript_name", "coder", "text", "memo", "weight",
        "anchor", "start", "end",
    ])
    for e in excerpts:
        writer.writerow([
            e.get("code_number", ""),
            e["code_name"],
            e.get("code_path", ""),
            e.get("transcript_label", ""),
            e.get("transcript_name", ""),
            e["coder"],
            e["text"],
            e.get("memo", ""),
            e.get("weight", ""),
            "yes" if e.get("anchor") else "",
            e.get("start", ""),
            e.get("end", ""),
        ])
    return buf.getvalue()


def export_analysis_docx(project, excerpts, mode="separate"):
    """Export excerpts as DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    doc.add_heading(f"Analys — {project['name']}", level=0)
    grouped = _group_excerpts(excerpts, mode)
    for g in grouped:
        prefix = f"{g['code_number']}. " if g["code_number"] else ""
        heading = doc.add_heading(f"{prefix}{g['code_path'] or g['code_name']}", level=1)
        try:
            color = g["code_color"].lstrip("#")
            r, gr, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(r, gr, b)
        except Exception:
            pass

        for e in g["excerpts"]:
            label = e.get("transcript_label", "")
            name = e.get("transcript_name", "")
            anchor = " 📌" if e.get("anchor") else ""
            p = doc.add_paragraph()
            run = p.add_run(f"{label}. {name} (kodare: {e['coder']}){anchor}")
            run.bold = True
            run.font.size = Pt(10)

            quote = doc.add_paragraph(e["text"])
            quote.style = "Quote" if "Quote" in [s.name for s in doc.styles] else None
            quote.paragraph_format.left_indent = Pt(18)

            if e.get("memo"):
                memo_p = doc.add_paragraph()
                memo_run = memo_p.add_run(f"Memo: {e['memo']}")
                memo_run.italic = True
                memo_run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_analysis_odt(project, excerpts, mode="separate"):
    """Export excerpts as ODT."""
    from odf.opendocument import OpenDocumentText
    from odf.text import P, H, Span
    from odf.style import Style, TextProperties, ParagraphProperties
    FO_NS = "urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"

    doc = OpenDocumentText()

    bold_style = Style(name="Bold", family="text")
    bp = TextProperties()
    bp.setAttrNS(FO_NS, "font-weight", "bold")
    bold_style.addElement(bp)
    doc.automaticstyles.addElement(bold_style)

    italic_style = Style(name="Italic", family="text")
    ip = TextProperties()
    ip.setAttrNS(FO_NS, "font-style", "italic")
    italic_style.addElement(ip)
    doc.automaticstyles.addElement(italic_style)

    indent_style = Style(name="Quote", family="paragraph")
    pp = ParagraphProperties()
    pp.setAttrNS(FO_NS, "margin-left", "1cm")
    indent_style.addElement(pp)
    doc.automaticstyles.addElement(indent_style)

    doc.text.addElement(H(outlinelevel=1, text=f"Analys — {project['name']}"))

    grouped = _group_excerpts(excerpts, mode)
    for g in grouped:
        prefix = f"{g['code_number']}. " if g["code_number"] else ""
        doc.text.addElement(H(outlinelevel=2, text=f"{prefix}{g['code_path'] or g['code_name']}"))

        for e in g["excerpts"]:
            label = e.get("transcript_label", "")
            name = e.get("transcript_name", "")
            anchor = " 📌" if e.get("anchor") else ""
            p = P()
            s = Span(stylename=bold_style)
            s.addText(f"{label}. {name} (kodare: {e['coder']}){anchor}")
            p.addElement(s)
            doc.text.addElement(p)

            q = P(stylename=indent_style)
            q.addText(e["text"])
            doc.text.addElement(q)

            if e.get("memo"):
                mp = P()
                ms = Span(stylename=italic_style)
                ms.addText(f"Memo: {e['memo']}")
                mp.addElement(ms)
                doc.text.addElement(mp)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
