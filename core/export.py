"""
export.py — Export coded data to CSV and Markdown.
"""
from __future__ import annotations

import csv
import io
from .annotation import load_all_coders
from .codebook import get_code, flat_list, build_tree
from .project import get_transcript_text


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

def export_csv(folder: str, project: dict, tid=None) -> str:
    """
    Export all annotations to CSV.
    If tid is given, only that transcript; otherwise all.
    Returns CSV as a string.
    """
    transcripts = project["transcripts"]
    if tid:
        transcripts = [t for t in transcripts if t["id"] == tid]

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["transcript", "coder", "code", "theme_path",
                     "start", "end", "text", "memo", "created"])

    for t in transcripts:
        all_coders = load_all_coders(folder, t["id"])
        for coder, anns in all_coders.items():
            for ann in anns:
                code = get_code(project, ann["code_id"])
                code_name = code["name"] if code else ann["code_id"]
                # Build breadcrumb path
                path = _code_path(project, ann["code_id"])
                writer.writerow([
                    t["name"], coder, code_name, path,
                    ann["start"], ann["end"], ann["text"],
                    ann.get("memo", ""), ann.get("created", ""),
                ])
    return output.getvalue()


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def export_markdown_by_code(folder: str, project: dict, tid=None) -> str:
    """
    Export quotes grouped by code/theme in Markdown format.
    """
    transcripts = project["transcripts"]
    if tid:
        transcripts = [t for t in transcripts if t["id"] == tid]

    # Gather all annotations keyed by code_id
    by_code: dict[str, list] = {}
    for t in transcripts:
        all_coders = load_all_coders(folder, t["id"])
        for coder, anns in all_coders.items():
            for ann in anns:
                cid = ann["code_id"]
                if cid not in by_code:
                    by_code[cid] = []
                by_code[cid].append({**ann, "transcript_name": t["name"], "coder": coder})

    if not by_code:
        return "_Inga kodningar hittades._\n"

    lines = [f"# {project['name']} — Kodade citat\n"]

    tree = build_tree(project)
    _render_tree_md(tree, by_code, lines, level=2)

    # Codes that exist in annotations but not in codebook
    known_ids = {c["id"] for c in project["codes"]}
    orphans = [cid for cid in by_code if cid not in known_ids]
    if orphans:
        lines.append("## Okända koder\n")
        for cid in orphans:
            lines.append(f"### {cid}\n")
            for ann in by_code[cid]:
                lines.append(_format_quote(ann))

    return "\n".join(lines)


def export_csv_tidy(folder: str, project: dict, tid=None) -> str:
    """
    Export annotations in tidy (long) format for R/Python analysis.
    One row per annotation with all metadata as columns.
    """
    transcripts = project["transcripts"]
    if tid:
        transcripts = [t for t in transcripts if t["id"] == tid]

    by_id = {c["id"]: c for c in project.get("codes", [])}
    proj_name = project.get("name", "")

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "project", "transcript", "transcript_category",
        "coder", "code", "parent_code", "code_path", "code_color",
        "start", "end", "text_length", "text", "memo",
        "weight", "anchor", "created",
    ])

    for t in transcripts:
        all_coders = load_all_coders(folder, t["id"])
        cat = t.get("category", "")
        for coder, anns in all_coders.items():
            for ann in anns:
                code = by_id.get(ann["code_id"])
                code_name   = code["name"] if code else ann["code_id"]
                parent_name = by_id.get(code["parent"], {}).get("name", "") if code and code.get("parent") else ""
                path        = _code_path(project, ann["code_id"])
                color       = code.get("color", "") if code else ""
                writer.writerow([
                    proj_name,
                    t["name"],
                    cat,
                    coder,
                    code_name,
                    parent_name,
                    path,
                    color,
                    ann["start"],
                    ann["end"],
                    ann["end"] - ann["start"],
                    ann["text"],
                    ann.get("memo", ""),
                    ann.get("weight", ""),
                    "TRUE" if ann.get("anchor") else "FALSE",
                    ann.get("created", ""),
                ])
    return output.getvalue()


def export_codebook_csv(project: dict, counts: dict | None = None) -> str:
    """
    Export the codebook as CSV.
    Columns: number, name, parent, description, count
    counts: optional {code_id: int} from stats.compute_stats
    """
    if counts is None:
        counts = {}

    tree = build_tree(project)
    _assign_numbers_py(tree, "")

    rows = []

    def _walk(nodes, parent_name):
        for node in nodes:
            rows.append({
                "number":      node.get("_number", ""),
                "name":        node["name"],
                "parent":      parent_name,
                "description": node.get("description", ""),
                "count":       counts.get(node["id"], 0),
            })
            _walk(node.get("children", []), node["name"])

    _walk(tree, "")

    out = io.StringIO()
    writer = csv.DictWriter(
        out,
        fieldnames=["number", "name", "parent", "description", "count"],
    )
    writer.writeheader()
    writer.writerows(rows)
    return out.getvalue()


def _assign_numbers_py(nodes: list, prefix: str):
    for i, node in enumerate(nodes):
        number = f"{prefix}.{i + 1}" if prefix else str(i + 1)
        node["_number"] = number
        _assign_numbers_py(node.get("children", []), number)


def export_markdown_codebook(project: dict) -> str:
    """Export the codebook as a Markdown document."""
    lines = [f"# Kodbok — {project['name']}\n"]
    tree = build_tree(project)
    _render_codebook_tree(tree, lines, level=2)
    if not tree:
        lines.append("_Kodboken är tom._\n")
    return "\n".join(lines)


def export_markdown_transcript(folder: str, project: dict, tid: str, coder: str) -> str:
    """Export a transcript with inline code annotations."""
    from .annotation import load_annotations
    t = next((t for t in project["transcripts"] if t["id"] == tid), None)
    if not t:
        return "_Transkript hittades inte._\n"

    text = get_transcript_text(folder, t)
    anns = load_annotations(folder, tid, coder)
    anns_sorted = sorted(anns, key=lambda a: a["start"])

    lines = [f"# {t['name']}\n", f"_Kodare: {coder}_\n\n---\n"]
    cursor = 0
    for ann in anns_sorted:
        s, e = ann["start"], ann["end"]
        if s < cursor:
            continue
        lines.append(text[cursor:s])
        code = get_code(project, ann["code_id"])
        code_name = code["name"] if code else ann["code_id"]
        lines.append(f"**[{code_name}]** *{text[s:e]}*")
        cursor = e
    lines.append(text[cursor:])
    lines.append("\n\n---\n\n## Kodsammanfattning\n")
    for ann in anns_sorted:
        code = get_code(project, ann["code_id"])
        code_name = code["name"] if code else ann["code_id"]
        lines.append(f"- **{code_name}**: {ann['text'][:80]}{'…' if len(ann['text']) > 80 else ''}")
        if ann.get("memo"):
            lines.append(f"  - _Memo: {ann['memo']}_")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _code_path(project: dict, code_id: str) -> str:
    """Return 'Grandparent > Parent > Code' breadcrumb."""
    by_id = {c["id"]: c for c in project["codes"]}
    chain = []
    cid = code_id
    while cid and cid in by_id:
        chain.insert(0, by_id[cid]["name"])
        cid = by_id[cid].get("parent")
    return " > ".join(chain)


def _format_quote(ann: dict) -> str:
    lines = [f"\n**{ann['transcript_name']}** _(kodare: {ann['coder']})_"]
    lines.append(f"> {ann['text']}")
    if ann.get("memo"):
        lines.append(f"> _Memo: {ann['memo']}_")
    return "\n".join(lines) + "\n"


def _render_tree_md(nodes: list, by_code: dict, lines: list, level: int):
    for node in nodes:
        hashes = "#" * level
        quotes = by_code.get(node["id"], [])
        if quotes or node["children"]:
            lines.append(f"{hashes} {node['name']}\n")
            if node.get("description"):
                lines.append(f"_{node['description']}_\n")
            for ann in quotes:
                lines.append(_format_quote(ann))
            _render_tree_md(node["children"], by_code, lines, level + 1)


def _render_codebook_tree(nodes: list, lines: list, level: int):
    for node in nodes:
        hashes = "#" * level
        lines.append(f"{hashes} {node['name']}\n")
        if node.get("description"):
            lines.append(f"{node['description']}\n")
        _render_codebook_tree(node["children"], lines, level + 1)


# ---------------------------------------------------------------------------
# Code tree DOCX / ODT exports
# ---------------------------------------------------------------------------

def export_codetree_docx(project: dict) -> bytes:
    """Export the code tree as a Word .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    import io

    tree = build_tree(project)
    doc = Document()
    doc.add_heading(f"Kodbok — {project['name']}", 0)

    def _hex_rgb(hex_color):
        h = hex_color.lstrip("#")
        if len(h) == 3:
            h = h[0]*2 + h[1]*2 + h[2]*2
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

    def _walk(nodes, depth):
        for node in nodes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(depth * 18)
            bullet = "■ " if depth == 0 else "▸ "
            run = p.add_run("  " * depth + bullet + node["name"])
            run.bold = (depth == 0)
            run.font.size = Pt(max(9, 13 - depth))
            try:
                r, g, b = _hex_rgb(node.get("color", "#888888"))
                run.font.color.rgb = RGBColor(r, g, b)
            except Exception:
                pass
            if node.get("description"):
                dp = doc.add_paragraph("  " * (depth + 1) + node["description"])
                dp.paragraph_format.left_indent = Pt(depth * 18 + 10)
                if dp.runs:
                    dp.runs[0].italic = True
                    dp.runs[0].font.size = Pt(10)
            _walk(node["children"], depth + 1)

    _walk(tree, 0)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_codetree_odt(project: dict) -> bytes:
    """Export the code tree as an OpenDocument Text .odt file."""
    try:
        from odf.opendocument import OpenDocumentText
        from odf.text import P, H
    except ImportError:
        raise RuntimeError("odfpy not installed. Run: pip install odfpy")
    import io

    tree = build_tree(project)
    doc = OpenDocumentText()

    title_el = H(outlinelevel=1)
    title_el.addText(f"Kodbok — {project['name']}")
    doc.text.addElement(title_el)

    def _walk(nodes, depth):
        for node in nodes:
            if depth == 0:
                h = H(outlinelevel=2)
                h.addText(node["name"])
                doc.text.addElement(h)
            else:
                p = P()
                p.addText("  " * depth + "• " + node["name"])
                doc.text.addElement(p)
            if node.get("description"):
                dp = P()
                dp.addText("  " * (depth + 1) + node["description"])
                doc.text.addElement(dp)
            _walk(node["children"], depth + 1)

    _walk(tree, 0)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
