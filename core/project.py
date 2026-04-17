"""
project.py — Project management for Transcribbler.
A project is a folder containing project.json, a transcripts/ subdir,
and an annotations/ subdir.
"""
from __future__ import annotations

import base64
import json
import os
import shutil
import uuid
from datetime import datetime
from pathlib import Path

import docx

PROJECT_FILE = "project.json"
TRANSCRIPTS_DIR = "transcripts"
ANNOTATIONS_DIR = "annotations"

# Importera bildstöd från ocr-modulen (undviker duplicering)
from core.ocr import SUPPORTED_IMAGES, is_image  # noqa: E402


def _now():
    return datetime.now().isoformat(timespec="seconds")


# ---------------------------------------------------------------------------
# Create / open
# ---------------------------------------------------------------------------

def create_project(folder: str, name: str, coder: str,
                   password: str | None = None) -> tuple:
    """Initialise a new project folder.

    Returns (project_dict, derived_key | None).
    """
    folder = Path(folder)
    folder.mkdir(parents=True, exist_ok=True)
    (folder / TRANSCRIPTS_DIR).mkdir(exist_ok=True)
    (folder / ANNOTATIONS_DIR).mkdir(exist_ok=True)

    project = {
        "name": name,
        "coder": coder,
        "created": _now(),
        "modified": _now(),
        "codes": [],
        "transcripts": [],
        "speakers": [],   # voice profiles (future feature)
        "numbering": True,
        "trans_order": True,
    }

    derived_key = None
    if password:
        from core.crypto import derive_key, make_verify_token
        salt = os.urandom(16)
        derived_key = derive_key(password, salt)
        verify_token = make_verify_token(derived_key)
        _save_encrypted(folder, project, derived_key, salt, verify_token)
    else:
        _save(folder, project)
    return project, derived_key


def open_project(folder: str, password: str | None = None) -> tuple:
    """Load an existing project.

    Returns (project_dict, derived_key | None).
    Raises ValueError if the project is encrypted and no/wrong password.
    """
    path = Path(folder) / PROJECT_FILE
    with open(path, encoding="utf-8") as f:
        raw = json.load(f)

    derived_key = None
    if raw.get("encrypted"):
        if not password:
            raise ValueError("encrypted")
        from core.crypto import derive_key, check_verify_token
        salt = base64.b64decode(raw["salt"])
        derived_key = derive_key(password, salt)
        token = base64.b64decode(raw["verify_token"])
        if not check_verify_token(token, derived_key):
            raise ValueError("wrong_password")
        from core.crypto import decrypt_blob
        payload = base64.b64decode(raw["payload"])
        project = json.loads(decrypt_blob(payload, derived_key))
    else:
        project = raw

    # Back-fill defaults for older projects missing these flags.
    if "numbering" not in project:
        project["numbering"] = True
    if "trans_order" not in project:
        project["trans_order"] = True
    return project, derived_key


def reload_project(folder: str, key: bytes | None = None) -> dict:
    """Re-read project.json using an existing derived key (no password needed).

    Used by background threads that need a fresh copy of the project dict.
    """
    path = Path(folder) / PROJECT_FILE
    with open(path, encoding="utf-8") as f:
        raw = json.load(f)

    if raw.get("encrypted"):
        if not key:
            raise ValueError("Project is encrypted but no key provided.")
        from core.crypto import decrypt_blob
        payload = base64.b64decode(raw["payload"])
        project = json.loads(decrypt_blob(payload, key))
    else:
        project = raw

    if "numbering" not in project:
        project["numbering"] = True
    if "trans_order" not in project:
        project["trans_order"] = True
    return project


def save_project(folder: str, project: dict, key: bytes | None = None):
    project["modified"] = _now()
    if key:
        _save_encrypted_update(Path(folder), project, key)
    else:
        _save(Path(folder), project)


def set_transcript_photos(folder: str, project: dict, tid: str, photos: list,
                          key: bytes | None = None) -> dict:
    """Attach a list of photo filenames to a transcript entry and save project.json."""
    for t in project.get("transcripts", []):
        if t["id"] == tid:
            t["photos"] = photos
            break
    save_project(folder, project, key)
    return project


def check_encrypted(folder: str) -> bool:
    """Return True if the project in *folder* is encrypted."""
    path = Path(folder) / PROJECT_FILE
    try:
        with open(path, encoding="utf-8") as f:
            raw = json.load(f)
        return bool(raw.get("encrypted"))
    except Exception:
        return False


def _save(folder: Path, project: dict):
    path = folder / PROJECT_FILE
    with open(path, "w", encoding="utf-8") as f:
        json.dump(project, f, ensure_ascii=False, indent=2)


def _save_encrypted(folder: Path, project: dict, key: bytes,
                    salt: bytes, verify_token: bytes):
    """Write project.json as an encrypted envelope (first save)."""
    from core.crypto import encrypt_blob
    payload = json.dumps(project, ensure_ascii=False, indent=2).encode("utf-8")
    envelope = {
        "encrypted": True,
        "salt": base64.b64encode(salt).decode(),
        "verify_token": base64.b64encode(verify_token).decode(),
        "payload": base64.b64encode(encrypt_blob(payload, key)).decode(),
    }
    path = folder / PROJECT_FILE
    with open(path, "w", encoding="utf-8") as f:
        json.dump(envelope, f, ensure_ascii=False, indent=2)


def _save_encrypted_update(folder: Path, project: dict, key: bytes):
    """Re-encrypt project.json preserving existing salt and verify_token."""
    path = folder / PROJECT_FILE
    with open(path, encoding="utf-8") as f:
        existing = json.load(f)
    from core.crypto import encrypt_blob
    payload = json.dumps(project, ensure_ascii=False, indent=2).encode("utf-8")
    existing["payload"] = base64.b64encode(encrypt_blob(payload, key)).decode()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# Transcripts
# ---------------------------------------------------------------------------

def add_transcript(folder: str, project: dict, src_path: str, name: str = "",
                   key: bytes | None = None) -> dict:
    """
    Copy a .txt or .docx file into the project's transcripts/ dir,
    extract text, register it in the project, return updated project.
    """
    src = Path(src_path)
    ext = src.suffix.lower()
    tid = str(uuid.uuid4())[:8]
    dest_name = f"{tid}{ext}"
    dest = Path(folder) / TRANSCRIPTS_DIR / dest_name

    shutil.copy2(src, dest)

    # For .md files: extract frontmatter metadata (title, category) before text extraction
    fm = _parse_md_frontmatter(src) if ext == ".md" else {}

    text, fmt_spans = _extract_text_with_formatting(src)
    # Save plain-text version alongside original
    txt_path = Path(folder) / TRANSCRIPTS_DIR / f"{tid}.txt"
    if key:
        from core.crypto import encrypt_text_file
        encrypt_text_file(txt_path, text, key)
    else:
        txt_path.write_text(text, encoding="utf-8")

    # Save formatting spans (bold/italic) if any were found in the document
    if fmt_spans:
        from core.formatting import save_formatting
        import uuid as _uuid
        enriched = []
        for s in fmt_spans:
            s["id"] = str(_uuid.uuid4())[:8]
            s["created"] = _now()
            enriched.append(s)
        save_formatting(folder, tid, "__import__", enriched, key=key)

    if name:
        display_name = name
    else:
        title = fm.get("title")
        date = str(fm.get("date", ""))[:10]
        if title and len(date) == 10 and date[4] == "-" and date[7] == "-":
            display_name = f"{title}_{date}"
        else:
            display_name = title or src.stem

    entry = {
        "id": tid,
        "name": display_name,
        "original": dest_name,
        "text_file": f"{tid}.txt",
        "tags": [],
        "added": _now(),
    }
    if fm.get("category"):
        entry["category"] = fm["category"]
    project["transcripts"].append(entry)
    save_project(folder, project, key)
    return project


def add_audio_transcript(folder: str, project: dict, tid: str, name: str,
                         audio_src_path: str, text: str,
                         segments: list, meta: dict,
                         key: bytes | None = None) -> dict:
    """
    Finalise an audio-sourced transcript after diarization + transcription.

    audio_src_path — path to the original audio temp file (will be copied)
    text           — speaker-labelled plain text
    segments       — [{speaker, start, end, text}]
    meta           — {whisper_model, language, diarization, diarization_settings,
                      speakers (name-mapping dict)}
    """
    import json as _json
    folder_path = Path(folder)
    audio_src = Path(audio_src_path)
    ext = audio_src.suffix.lower()

    # Copy audio file into project
    audio_dest_name = f"{tid}{ext}"
    audio_dest = folder_path / TRANSCRIPTS_DIR / audio_dest_name
    shutil.copy2(audio_src, audio_dest)

    # Write extracted plain text
    txt_path = folder_path / TRANSCRIPTS_DIR / f"{tid}.txt"
    if key:
        from core.crypto import encrypt_text_file, encrypt_json_file
        encrypt_text_file(txt_path, text, key)
    else:
        txt_path.write_text(text, encoding="utf-8")

    # Write segments (separate file — can be large)
    seg_path = folder_path / TRANSCRIPTS_DIR / f"{tid}_segments.json"
    if key:
        from core.crypto import encrypt_json_file as _ejf
        _ejf(seg_path, segments, key)
    else:
        seg_path.write_text(
            _json.dumps(segments, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    entry = {
        "id": tid,
        "name": name,
        "source": "audio",
        "original": audio_dest_name,
        "text_file": f"{tid}.txt",
        "audio_file": audio_dest_name,
        "whisper_model": meta.get("whisper_model", "medium"),
        "language": meta.get("language", "sv"),
        "diarization": meta.get("diarization", False),
        "diarization_settings": meta.get("diarization_settings", {}),
        "speakers": meta.get("speakers", {}),
        "tags": [],
        "added": _now(),
    }
    project["transcripts"].append(entry)
    save_project(folder, project, key)
    return project


def add_image_transcript(folder: str, project: dict, tid: str, name: str,
                         image_src_path: str, text: str,
                         key: bytes | None = None) -> dict:
    """
    Finalise an image-sourced transcript after OCR.

    image_src_path — path to the (temp) image file (will be copied into project)
    text           — OCR-extracted plain text
    """
    folder_path = Path(folder)
    image_src = Path(image_src_path)
    ext = image_src.suffix.lower()

    # Copy source image into project
    source_dest_name = f"{tid}_source{ext}"
    source_dest = folder_path / TRANSCRIPTS_DIR / source_dest_name
    shutil.copy2(image_src, source_dest)

    # Write extracted plain text
    txt_path = folder_path / TRANSCRIPTS_DIR / f"{tid}.txt"
    if key:
        from core.crypto import encrypt_text_file
        encrypt_text_file(txt_path, text, key)
    else:
        txt_path.write_text(text, encoding="utf-8")

    entry = {
        "id": tid,
        "name": name,
        "source": "image",
        "source_file": source_dest_name,
        "text_file": f"{tid}.txt",
        "tags": [],
        "added": _now(),
    }
    project["transcripts"].append(entry)
    save_project(folder, project, key)
    return project


def _safe_child(base: Path, name: str) -> Path | None:
    # Resolve `base / name` and ensure it stays under `base`.
    # Returns None if `name` contains path-traversal (e.g. "../etc/passwd")
    # or is absolute. Protects routes that consume user-controlled file
    # names from project.json.
    base = Path(base).resolve()
    candidate = (base / name).resolve()
    try:
        candidate.relative_to(base)
    except ValueError:
        return None
    return candidate


def get_transcript_text(folder: str, transcript: dict,
                        key: bytes | None = None) -> str:
    tdir = Path(folder) / TRANSCRIPTS_DIR
    path = _safe_child(tdir, transcript["text_file"])
    if path is None:
        raise ValueError("invalid text_file path")
    if key:
        from core.crypto import is_encrypted_file, decrypt_text_file
        if is_encrypted_file(path):
            return decrypt_text_file(path, key)
    return path.read_text(encoding="utf-8")


def remove_transcript(folder: str, project: dict, tid: str,
                      key: bytes | None = None) -> dict:
    t = next((t for t in project["transcripts"] if t["id"] == tid), None)
    if not t:
        return project
    tdir = Path(folder) / TRANSCRIPTS_DIR
    for fname in [t.get("original"), t.get("text_file"), t.get("source_file")]:
        if fname:
            p = _safe_child(tdir, fname)
            if p and p.exists():
                p.unlink()
    # Remove attached photos (imported from Notescribbler)
    for fname in t.get("photos", []):
        if fname:
            p = _safe_child(tdir, fname)
            if p and p.exists():
                p.unlink()
    # Remove segments file (audio transcripts)
    seg_path = _safe_child(tdir, f"{tid}_segments.json")
    if seg_path and seg_path.exists():
        seg_path.unlink()
    # Remove OCR boxes file (image transcripts)
    boxes_path = _safe_child(tdir, f"{tid}_ocr_boxes.json")
    if boxes_path and boxes_path.exists():
        boxes_path.unlink()
    project["transcripts"] = [t for t in project["transcripts"] if t["id"] != tid]
    # Remove all annotation files for this transcript
    ann_dir = Path(folder) / ANNOTATIONS_DIR
    ann_dir_resolved = ann_dir.resolve()
    for f in ann_dir.glob(f"{tid}.*.json"):
        try:
            f.resolve().relative_to(ann_dir_resolved)
        except ValueError:
            continue
        f.unlink()
    save_project(folder, project, key)
    return project


def _parse_md_frontmatter(path: Path) -> dict:
    """
    Extract top-level scalar fields from YAML frontmatter in a .md file.
    Returns a dict with string values for keys found (e.g. title, category).
    List fields (tags, photos) and nested objects (location) are ignored.
    Returns {} if no frontmatter present.
    """
    import re
    raw = path.read_text(encoding="utf-8", errors="replace")
    m = re.match(r"\A---\n(.*?)\n---\n?", raw, re.DOTALL)
    if not m:
        return {}
    result = {}
    for line in m.group(1).splitlines():
        kv = re.match(r"^([a-zA-Z_]\w*):\s*(.+)", line)
        if not kv:
            continue
        key, value = kv.group(1), kv.group(2).strip()
        # Skip list values and nested objects
        if value.startswith("[") or value.startswith("{"):
            continue
        # Strip surrounding quotes
        if len(value) >= 2 and value[0] in ('"', "'") and value[-1] == value[0]:
            value = value[1:-1]
        result[key] = value
    return result


def _extract_text(path: Path) -> str:
    text, _ = _extract_text_with_formatting(path)
    return text


def _extract_text_with_formatting(path: Path) -> tuple:
    """Return (plain_text, formatting_spans) where spans are [{start, end, type}]."""
    ext = path.suffix.lower()
    if ext == ".docx":
        return _extract_docx_with_formatting(path)
    elif ext == ".odt":
        return _extract_odt_with_formatting(path)
    elif ext == ".md":
        import re
        raw = path.read_text(encoding="utf-8", errors="replace")
        raw = re.sub(r"\A---\n.*?\n---\n?", "", raw, count=1, flags=re.DOTALL)
        raw = re.sub(r"^#{1,6}\s+", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", raw)
        raw = re.sub(r"_{1,2}(.+?)_{1,2}", r"\1", raw)
        raw = re.sub(r"`{1,3}[^`]*`{1,3}", "", raw)
        raw = re.sub(r"!\[.*?\]\(.*?\)", "", raw)
        raw = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", raw)
        return raw, []
    else:
        return _read_text_autodetect(path), []


def _extract_docx_with_formatting(path: Path) -> tuple:
    """Extract text and bold/italic spans from a .docx file."""
    doc = docx.Document(str(path))
    lines = []
    spans = []
    offset = 0
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            start = offset
            end = offset + len(text)
            if run.bold:
                spans.append({"start": start, "end": end, "type": "bold"})
            if run.italic:
                spans.append({"start": start, "end": end, "type": "italic"})
            offset = end
        lines.append(offset)
        offset += 1  # newline
    plain = "\n".join(p.text for p in doc.paragraphs)
    return plain, spans


def _extract_odt_with_formatting(path: Path) -> tuple:
    """Extract text and bold/italic spans from an .odt file."""
    from odf import text as odftext, teletype
    from odf.opendocument import load as odf_load
    from odf.style import TextProperties
    FO_NS = "urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"
    doc = odf_load(str(path))
    paras = doc.text.getElementsByType(odftext.P)
    styles = {}
    for s in doc.automaticstyles.childNodes:
        name = s.getAttribute("name")
        if name:
            tp = s.getElementsByType(TextProperties)
            if tp:
                props = tp[0]
                fw = props.getAttrNS(FO_NS, "font-weight") or ""
                fs = props.getAttrNS(FO_NS, "font-style") or ""
                styles[name] = {
                    "bold": fw == "bold",
                    "italic": fs == "italic",
                }
    plain_parts = []
    spans = []
    offset = 0
    for para in paras:
        para_text = teletype.extractText(para)
        child_offset = offset
        for child in para.childNodes:
            qname = getattr(child, "qname", None)
            if qname is None:
                child_text = str(child)
            else:
                child_text = teletype.extractText(child)
            if not child_text:
                continue
            start = child_offset
            end = child_offset + len(child_text)
            if qname and qname[1] == "span":
                style_name = child.getAttribute("stylename")
                if style_name and style_name in styles:
                    s = styles[style_name]
                    if s.get("bold"):
                        spans.append({"start": start, "end": end, "type": "bold"})
                    if s.get("italic"):
                        spans.append({"start": start, "end": end, "type": "italic"})
            child_offset = end
        plain_parts.append(para_text)
        offset += len(para_text) + 1
    return "\n".join(plain_parts), spans


def _read_text_autodetect(path: Path) -> str:
    """Read a plain-text file, detecting the encoding when possible."""
    raw = path.read_bytes()
    try:
        from charset_normalizer import from_bytes
        best = from_bytes(raw).best()
        if best is not None:
            return str(best)
    except Exception:
        pass
    return raw.decode("utf-8", errors="replace")
